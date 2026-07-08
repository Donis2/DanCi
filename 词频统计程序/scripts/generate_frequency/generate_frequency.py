# -*- coding: utf-8 -*-
"""
词频统计脚本

输入:
  - 词汇/ 考研英语大纲 docx (提取词表 + 释义)
  - 真题/ 所有 docx (提取试卷文本,统计词频)

输出:
  - netem_full_list.json  (覆盖原文件,原文件备份为 .bak)
  - netem_full_list.sql

策略:
  1. 解析大纲,得到 headword -> 释义 映射
  2. 遍历所有真题 docx,提取段落 + 表格单元格文本
  3. 用正则提取所有英文 token,小写化
  4. 用 nltk WordNetLemmatizer 做词形还原 (依次尝试 v/n/a/r,优先匹配词表)
  5. 只统计 lemma 命中词表的 token
  6. 从原 netem_full_list.json 继承 其他拼写/分类/子分类 (按 word 匹配)
  7. 按词频降序排序,输出 json 和 sql
"""

import os
import re
import json
import shutil
from collections import Counter
from docx import Document
from nltk.stem import WordNetLemmatizer

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
VOCAB_DIR = os.path.join(BASE, "词汇")
EXAM_DIR = os.path.join(BASE, "真题")
OUT_JSON = os.path.join(BASE, "netem_full_list.json")
OUT_SQL = os.path.join(BASE, "netem_full_list.sql")

# 词性标记正则 (用于定位词条锚点)
POS_PATTERN = r'(?:n|v|a|ad|vt|vi|art|prep|conj|pron|num|int|aux|mod|abbr|pl)\.(?:&[a-z]+\.?)*'
ENTRY_RE = re.compile(r'([a-zA-Z][a-zA-Z\'-]*)\s+(' + POS_PATTERN + r')')

# 英文 token 正则 (不含撇号: 所有格 people's 会拆成 people + s, s 不在词表会被跳过)
TOKEN_RE = re.compile(r"[a-zA-Z][a-zA-Z-]*")

# 英文缩写展开 (小写),用于还原 is/are/am/have/not/will 等
CONTRACTIONS_MAP = {
    "it's": "it is", "that's": "that is", "he's": "he is", "she's": "she is",
    "what's": "what is", "there's": "there is", "here's": "here is",
    "who's": "who is", "where's": "where is", "how's": "how is",
    "they're": "they are", "we're": "we are", "you're": "you are",
    "who're": "who are", "what're": "what are", "where're": "where are",
    "i'm": "i am",
    "they've": "they have", "we've": "we have", "you've": "you have",
    "i've": "i have", "could've": "could have", "would've": "would have",
    "should've": "should have", "might've": "might have", "must've": "must have",
    "they'll": "they will", "we'll": "we will", "you'll": "you will",
    "i'll": "i will", "he'll": "he will", "she'll": "she will",
    "it'll": "it will", "that'll": "that will", "who'll": "who will",
    "don't": "do not", "doesn't": "does not", "didn't": "did not",
    "isn't": "is not", "aren't": "are not", "wasn't": "was not",
    "weren't": "were not", "couldn't": "could not", "wouldn't": "would not",
    "shouldn't": "should not", "haven't": "have not", "hasn't": "has not",
    "hadn't": "had not", "can't": "can not", "won't": "will not",
    "shan't": "shall not", "needn't": "need not", "oughtn't": "ought not",
    "mightn't": "might not", "mustn't": "must not",
}
CONTRACTIONS_RE = re.compile(
    r'\b(?:' + '|'.join(re.escape(k) for k in CONTRACTIONS_MAP) + r')\b'
)


def expand_contractions(text):
    """展开英文缩写 (文本应已小写)"""
    return CONTRACTIONS_RE.sub(lambda m: CONTRACTIONS_MAP[m.group(0)], text)


def parse_vocab_docx(path):
    """解析词汇大纲 docx,返回 [(headword, definition), ...]
    headword 保留原大小写 (如 April, I, Bible)
    """
    doc = Document(path)
    entries = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # 找所有 "word POS." 锚点
        matches = list(ENTRY_RE.finditer(text))
        if not matches:
            continue
        for i, m in enumerate(matches):
            head = m.group(1)  # 保留原大小写
            entry_start = m.start(1)
            if i + 1 < len(matches):
                entry_end = matches[i + 1].start(1)
            else:
                entry_end = len(text)
            entry_text = text[entry_start:entry_end].strip()
            # 去掉结尾残留的数字序号
            entry_text = re.sub(r'\s*\d+\.\s*$', '', entry_text).strip()
            entries.append((head, entry_text))
    return entries


def build_vocab_map(entries):
    """构建 headword -> definition 映射 (重复词取首个,大小写不敏感去重)"""
    vocab = {}
    seen_lower = set()
    for head, defn in entries:
        low = head.lower()
        if low not in seen_lower:
            seen_lower.add(low)
            vocab[head] = defn
    return vocab


def build_lower_to_head(vocab):
    """构建 lowercase -> original_headword 映射"""
    return {head.lower(): head for head in vocab}


def extract_exam_text(path):
    """提取一份真题 docx 的所有文本 (段落 + 表格单元格)"""
    doc = Document(path)
    chunks = []
    for p in doc.paragraphs:
        t = p.text
        if t:
            chunks.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                if t:
                    chunks.append(t)
    return "\n".join(chunks)


def iter_exam_docxs(root):
    """递归遍历真题文件夹下所有 docx"""
    for dirpath, _, filenames in os.walk(root):
        for fn in filenames:
            if fn.lower().endswith(".docx"):
                yield os.path.join(dirpath, fn)


def normalize_token(token, lower_to_head, lemmatizer):
    """词形还原: 返回词表中对应的原大小写 headword,或 None
    token 已小写。依次尝试: 直接匹配 -> v/n/a/r lemmatize 后匹配
    """
    # 1. 直接命中词表 (含 KEEP_AS_IS 的词也在词表里)
    if token in lower_to_head:
        return lower_to_head[token]
    # 2. lemmatize 后匹配
    for pos in ('v', 'n', 'a', 'r'):
        lemma = lemmatizer.lemmatize(token, pos)
        if lemma != token and lemma in lower_to_head:
            return lower_to_head[lemma]
    return None


def load_original_meta(path):
    """从原 json 读取 word -> (variant, category, subcategory) 映射"""
    if not os.path.exists(path):
        return {}
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    key = list(data.keys())[0]
    meta = {}
    for item in data[key]:
        meta[item["单词"]] = (
            item.get("其他拼写"),
            item.get("分类"),
            item.get("子分类"),
        )
    return meta


def main():
    # 1. 解析大纲
    vocab_files = [f for f in os.listdir(VOCAB_DIR) if f.lower().endswith(".docx")]
    if not vocab_files:
        print("词汇目录下未找到 docx 文件")
        return
    vocab_path = os.path.join(VOCAB_DIR, vocab_files[0])
    print(f"[1/5] 解析词汇大纲: {vocab_path}")
    entries = parse_vocab_docx(vocab_path)
    vocab = build_vocab_map(entries)
    print(f"      共解析出 {len(vocab)} 个词条")

    # 2. 提取真题文本
    print(f"[2/5] 提取真题文本: {EXAM_DIR}")
    all_tokens = []
    exam_count = 0
    for docx_path in iter_exam_docxs(EXAM_DIR):
        try:
            text = extract_exam_text(docx_path)
            text = text.replace("\u2019", "'")  # 智能撇号转 ASCII 撇号
            text = expand_contractions(text.lower())
            tokens = TOKEN_RE.findall(text)
            all_tokens.extend(tokens)
            exam_count += 1
        except Exception as e:
            print(f"      [警告] 读取失败 {os.path.basename(docx_path)}: {e}")
    print(f"      共读取 {exam_count} 份真题, {len(all_tokens)} 个原始 token")

    # 3. 词形还原 + 统计 (只统计命中词表的)
    print("[3/5] 词形还原并统计词频")
    lemmatizer = WordNetLemmatizer()
    lower_to_head = build_lower_to_head(vocab)
    raw_counter = Counter(all_tokens)
    freq = Counter()  # key = 原大小写 headword
    hit = 0
    skipped = 0
    for token, cnt in raw_counter.items():
        headword = normalize_token(token, lower_to_head, lemmatizer)
        if headword is not None:
            freq[headword] += cnt
            hit += 1
        else:
            skipped += 1
    print(f"      unique token: {len(raw_counter)}, 命中词表: {hit}, 跳过: {skipped}")

    # 4. 继承原数据中的 variant/category/subcategory
    print("[4/5] 合并原数据中的 其他拼写/分类/子分类")
    # 优先从 .bak 读取原始 meta (避免二次运行时读到已覆盖数据)
    meta_path = OUT_JSON + ".bak" if os.path.exists(OUT_JSON + ".bak") else OUT_JSON
    meta = load_original_meta(meta_path)

    # 5. 组装结果 (词频降序,同频按字母序)
    rows = []
    for word, defn in vocab.items():
        count = freq.get(word, 0)
        variant, category, subcategory = meta.get(word, (None, None, None))
        rows.append({
            "word": word,
            "frequency": count,
            "definition": defn,
            "variant": variant,
            "category": category,
            "subcategory": subcategory,
        })
    # 排序: 词频降序, 同频按单词字母升序
    rows.sort(key=lambda r: (-r["frequency"], r["word"]))

    # 组装最终结构
    result_list = []
    for idx, r in enumerate(rows, 1):
        result_list.append({
            "序号": idx,
            "词频": r["frequency"],
            "单词": r["word"],
            "释义": r["definition"],
            "其他拼写": r["variant"],
            "分类": r["category"],
            "子分类": r["subcategory"],
        })
    title = f"{len(vocab)}考研词汇词频排序表"
    result = {title: result_list}

    # 备份原文件 (仅在 .bak 不存在时备份,避免覆盖原始备份)
    if os.path.exists(OUT_JSON) and not os.path.exists(OUT_JSON + ".bak"):
        shutil.copy2(OUT_JSON, OUT_JSON + ".bak")
        print(f"      已备份原 json -> {OUT_JSON}.bak")

    # 写 json
    with open(OUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"[5/5] 已写入 {OUT_JSON}")

    # 写 sql
    write_sql(OUT_SQL, title, result_list)
    print(f"      已写入 {OUT_SQL}")

    # 简要统计
    nonzero = sum(1 for r in result_list if r["词频"] > 0)
    top5 = result_list[:5]
    print("\n=== 统计摘要 ===")
    print(f"大纲词数: {len(vocab)}")
    print(f"试卷中出现过的词: {nonzero}")
    print(f"未出现的词: {len(vocab) - nonzero}")
    print("Top 5:")
    for r in top5:
        print(f"  {r['序号']:>4}. {r['单词']:<15} {r['词频']}")


def write_sql(path, table_title, rows):
    """写 sql 文件,结构与原 netem_full_list.sql 对齐"""
    lines = [
        "SET NAMES utf8mb4;",
        "SET FOREIGN_KEY_CHECKS = 0;",
        "",
        "-- ----------------------------",
        f"-- Table structure for netem_full_list",
        "-- ----------------------------",
        "DROP TABLE IF EXISTS `netem_full_list`;",
        "CREATE TABLE `netem_full_list`  (",
        "  `id` int(0) NOT NULL,",
        "  `frequency` int(0) NULL DEFAULT NULL,",
        "  `word` varchar(255) CHARACTER SET utf8mb4 COLLATE utf8mb4_bin NULL DEFAULT NULL,",
        "  `definition` varchar(255) CHARACTER SET utf8mb4 COLLATE utf8mb4_bin NULL DEFAULT NULL,",
        "  `variant` varchar(255) CHARACTER SET utf8mb4 COLLATE utf8mb4_bin NULL DEFAULT NULL,",
        "  `category` varchar(255) CHARACTER SET utf8mb4 COLLATE utf8mb4_bin NULL DEFAULT NULL,",
        "  `subcategory` varchar(255) CHARACTER SET utf8mb4 COLLATE utf8mb4_bin NULL DEFAULT NULL,",
        "  PRIMARY KEY (`id`) USING BTREE",
        ") ENGINE = InnoDB CHARACTER SET = utf8mb4 COLLATE = utf8mb4_bin ROW_FORMAT = Dynamic;",
        "",
        "-- ----------------------------",
        "-- Records of netem_full_list",
        "-- ----------------------------",
        "",
    ]

    def sql_str(v):
        if v is None:
            return "NULL"
        return "'" + str(v).replace("'", "''") + "'"

    for r in rows:
        values = [
            str(r["序号"]),
            str(r["词频"]),
            sql_str(r["单词"]),
            sql_str(r["释义"]),
            sql_str(r["其他拼写"]),
            sql_str(r["分类"]),
            sql_str(r["子分类"]),
        ]
        lines.append(f"INSERT INTO `netem_full_list` VALUES ({', '.join(values)});")
        lines.append("")

    lines.append("SET FOREIGN_KEY_CHECKS = 1;")

    with open(path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))


if __name__ == "__main__":
    main()
